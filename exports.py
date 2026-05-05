"""
exports.py — Generación de reportes en PDF.
"""

import os
import shutil
import datetime
import sqlite3

from copy import copy as _copy
from tkinter import messagebox, filedialog

from config import DB_PATH, REPORTES_DIR
from db import config_get


# ─────────────────────────────────────────────
#  UTILIDAD: NOMBRE DE ARCHIVO SEGURO
# ─────────────────────────────────────────────
def _safe_filename(name: str) -> str:
    import unicodedata
    nfkd = unicodedata.normalize("NFKD", name)
    ascii_name = nfkd.encode("ascii", "ignore").decode("ascii")
    for ch in r'\/:*?"<>|':
        ascii_name = ascii_name.replace(ch, "_")
    ascii_name = " ".join(ascii_name.split())
    ascii_name = ascii_name.replace(" ", "_")
    return ascii_name or "reporte"


def _pedir_ruta_guardado(nombre_sugerido: str) -> str | None:
    ruta = filedialog.asksaveasfilename(
        title="Guardar reporte PDF",
        initialdir=REPORTES_DIR,
        initialfile=nombre_sugerido,
        defaultextension=".pdf",
        filetypes=[("Archivo PDF", "*.pdf"), ("Todos los archivos", "*.*")],
    )
    return ruta if ruta else None


# ─────────────────────────────────────────────
#  CONVERSIÓN DOCX → PDF  (3 intentos)
# ─────────────────────────────────────────────
def _convertir_docx_a_pdf(fname_docx: str, fname_pdf: str) -> bool:
    """
    Intenta convertir el .docx al .pdf por tres métodos distintos.
    Retorna True si tuvo éxito.
    """

    # ── Intento 1: docx2pdf ──────────────────
    try:
        from docx2pdf import convert as _d2p_convert
        _d2p_convert(fname_docx, fname_pdf)
        if os.path.isfile(fname_pdf):
            print(f"✅ Convertido con docx2pdf → {fname_pdf}")
            return True
    except Exception as e:
        print(f"⚠ docx2pdf falló: {e}")

    # ── Intento 2: Word COM directo ──────────
    try:
        import comtypes.client          # noqa: F401  (importado dinámicamente)
        import comtypes                 # noqa: F401

        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc_com = word.Documents.Open(os.path.abspath(fname_docx))
        doc_com.SaveAs(os.path.abspath(fname_pdf), FileFormat=17)   # 17 = wdFormatPDF
        doc_com.Close()
        word.Quit()
        if os.path.isfile(fname_pdf):
            print(f"✅ Convertido con Word COM → {fname_pdf}")
            return True
    except ImportError:
        print("⚠ comtypes no disponible — omitiendo intento Word COM.")
    except Exception as e:
        print(f"⚠ Word COM falló: {e}")

    # ── Intento 3: LibreOffice ───────────────
    try:
        import subprocess
        libreoffice_bins = [
            "libreoffice",
            "soffice",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        pdf_dir  = os.path.dirname(fname_pdf)
        pdf_base = os.path.splitext(os.path.basename(fname_docx))[0] + ".pdf"
        pdf_temp = os.path.join(pdf_dir, pdf_base)

        for bin_path in libreoffice_bins:
            try:
                res = subprocess.run(
                    [bin_path, "--headless", "--convert-to", "pdf",
                        "--outdir", pdf_dir, fname_docx],
                    capture_output=True, timeout=60)
                if res.returncode == 0 and os.path.isfile(pdf_temp):
                    if pdf_temp != fname_pdf:
                        os.replace(pdf_temp, fname_pdf)
                    print(f"✅ Convertido con LibreOffice → {fname_pdf}")
                    return True
            except FileNotFoundError:
                continue
            except Exception as e:
                print(f"⚠ LibreOffice '{bin_path}' falló: {e}")
                continue
    except Exception as e:
        print(f"⚠ Error en intento LibreOffice: {e}")

    return False


# ─────────────────────────────────────────────
#  EXPORTAR PDF — DESDE PLANTILLA .docx
# ─────────────────────────────────────────────
def exportar_pdf_desde_docx(capacitacion_id: int, nombre_cap: str,
                            ruta_plantilla: str) -> str | None:
    try:
        import docx as _docx_mod       # noqa: F401
    except ImportError:
        messagebox.showerror(
            "Librería faltante",
            "Instala python-docx:\n\npip install python-docx")
        return None

    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.shared import Inches
    from copy import deepcopy

    with sqlite3.connect(DB_PATH) as con:
        rows = con.execute("""
            SELECT p.nombre || ' ' || p.apellido, p.cargo, a.hora_registro
            FROM asistencias a JOIN personas p ON p.id = a.persona_id
            WHERE a.capacitacion_id=? ORDER BY a.hora_registro
        """, (capacitacion_id,)).fetchall()
        cap_info = con.execute(
            "SELECT nombre, descripcion, fecha, firma_responsable, firma_png "
            "FROM capacitaciones WHERE id=?",
            (capacitacion_id,)).fetchone()

    hoy = datetime.date.today().strftime("%d/%m/%Y")
    firma_png_ruta = cap_info[4] if cap_info and len(cap_info) > 4 else None

    marcadores_globales = {
        "cap_nombre":        cap_info[0] if cap_info else "",
        "cap_descripcion":   cap_info[1] if cap_info and cap_info[1] else "",
        "cap_fecha":         cap_info[2] if cap_info else "",
        "cap_total":         str(len(rows)),
        "total":             str(len(rows)),
        "fecha_hoy":         hoy,
        "capacitacion":      cap_info[0] if cap_info else "",
        "fecha_cap":         cap_info[2] if cap_info else "",
        "firma_responsable": cap_info[3] if cap_info and cap_info[3] else "",
    }
    marcadores_fila_keys = {"nombre_completo", "cargo", "hora_registro", "numero"}

    nombre_sugerido = f"Asistencias_{_safe_filename(nombre_cap)}_{datetime.date.today()}.pdf"
    fname_pdf = _pedir_ruta_guardado(nombre_sugerido)
    if fname_pdf is None:
        return None

    import tempfile
    _tmp_dir = tempfile.mkdtemp()
    fname_docx = os.path.join(
        _tmp_dir,
        f"Asistencias_{_safe_filename(nombre_cap)}_{datetime.date.today()}.docx")
    shutil.copy2(ruta_plantilla, fname_docx)
    doc = DocxDocument(fname_docx)

    def _reemplazar_en_xml(elemento, mapa):
        for para_xml in elemento.iter(qn("w:p")):
            runs = para_xml.findall(".//" + qn("w:r"))
            if not runs:
                continue
            textos = []
            primer_t = None
            for r in runs:
                for t in r.findall(qn("w:t")):
                    textos.append(t.text or "")
                    if primer_t is None:
                        primer_t = t
                    else:
                        t.text = ""
            if primer_t is not None:
                texto_unido = "".join(textos)
                for k, v in mapa.items():
                    texto_unido = texto_unido.replace(f"{{{{{k}}}}}", str(v))
                primer_t.text = texto_unido
                if texto_unido != texto_unido.strip():
                    primer_t.set(
                        "{http://www.w3.org/XML/1998/namespace}space", "preserve")

    def _insertar_firma_imagen_en_parrafos(contenedor):
        if not firma_png_ruta or not os.path.isfile(firma_png_ruta):
            return
        for para in contenedor.paragraphs:
            texto_completo = "".join(r.text for r in para.runs)
            if "{{firma_imagen}}" not in texto_completo:
                continue
            for run in para.runs:
                run.text = ""
            run_img = para.add_run()
            run_img.add_picture(firma_png_ruta, width=Inches(1.5))

    def _insertar_firma_imagen_en_tablas(tabla):
        if not firma_png_ruta or not os.path.isfile(firma_png_ruta):
            return
        for fila in tabla.rows:
            for celda in fila.cells:
                for para in celda.paragraphs:
                    texto_completo = "".join(r.text for r in para.runs)
                    if "{{firma_imagen}}" not in texto_completo:
                        continue
                    for run in para.runs:
                        run.text = ""
                    run_img = para.add_run()
                    run_img.add_picture(firma_png_ruta, width=Inches(1.5))

    def _celda_tiene_marcador_fila(celda):
        return any(f"{{{{{mk}}}}}" in celda.text for mk in marcadores_fila_keys)

    def _fila_tiene_marcador_fila(fila):
        return any(_celda_tiene_marcador_fila(c) for c in fila.cells)

    for tabla in doc.tables:
        fila_template_idx = None
        for i, fila in enumerate(tabla.rows):
            if _fila_tiene_marcador_fila(fila):
                fila_template_idx = i
                break

        if fila_template_idx is not None:
            fila_tpl  = tabla.rows[fila_template_idx]
            tpl_xml   = deepcopy(fila_tpl._tr)
            tbl_xml   = tabla._tbl
            tbl_xml.remove(fila_tpl._tr)
            todas_filas  = tbl_xml.findall(qn("w:tr"))
            insert_after = todas_filas[fila_template_idx - 1] if fila_template_idx > 0 else None

            for idx, (nombre_completo, cargo, hora) in enumerate(rows):
                nueva_fila_xml = deepcopy(tpl_xml)
                mapa_fila = {
                    "numero":          str(idx + 1),
                    "nombre_completo": nombre_completo,
                    "cargo":           cargo or "",
                    "hora_registro":   hora,
                    **marcadores_globales,
                }
                _reemplazar_en_xml(nueva_fila_xml, mapa_fila)
                if insert_after is not None:
                    insert_after.addnext(nueva_fila_xml)
                    insert_after = nueva_fila_xml
                else:
                    tbl_xml.insert(0, nueva_fila_xml)
                    insert_after = nueva_fila_xml
        else:
            _reemplazar_en_xml(tabla._tbl, marcadores_globales)
        _insertar_firma_imagen_en_tablas(tabla)

    for para in doc.paragraphs:
        _reemplazar_en_xml(para._p, marcadores_globales)
    _insertar_firma_imagen_en_parrafos(doc)

    for section in doc.sections:
        for container in (section.header, section.footer):
            _reemplazar_en_xml(container._element, marcadores_globales)
            _insertar_firma_imagen_en_parrafos(container)

    doc.save(fname_docx)

    # ── Convertir a PDF ──────────────────────
    convertido = _convertir_docx_a_pdf(fname_docx, fname_pdf)

    if convertido and os.path.isfile(fname_pdf):
        shutil.rmtree(_tmp_dir, ignore_errors=True)
        return fname_pdf

    # Fallback: entregar el .docx si no se pudo convertir
    fname_docx_dest = os.path.splitext(fname_pdf)[0] + ".docx"
    shutil.copy2(fname_docx, fname_docx_dest)
    shutil.rmtree(_tmp_dir, ignore_errors=True)
    messagebox.showwarning(
        "Conversión a PDF no disponible",
        f"El documento Word se guardó en:\n{fname_docx_dest}\n\n"
        f"No se encontró Word ni LibreOffice para convertir a PDF.\n\n"
        f"Opciones:\n"
        f"  • Instala LibreOffice (gratis) y vuelve a exportar.\n"
        f"  • Abre el .docx en Word y guárdalo como PDF manualmente.\n"
        f"  • Ejecuta:  pip install docx2pdf")
    return fname_docx_dest


# ─────────────────────────────────────────────
#  EXPORTAR PDF — FORMATO CORPORATIVO (ReportLab)
# ─────────────────────────────────────────────
def exportar_pdf(capacitacion_id: int, nombre_cap: str) -> str | None:
    """
    Exporta a PDF.
    - Si hay plantilla .docx configurada: la rellena y convierte a PDF.
    - Si no: genera un PDF corporativo con ReportLab.
    """
    ruta_plantilla_docx = config_get("plantilla_pdf_docx")
    if ruta_plantilla_docx and os.path.isfile(ruta_plantilla_docx):
        return exportar_pdf_desde_docx(capacitacion_id, nombre_cap, ruta_plantilla_docx)

    # ── Fallback: PDF corporativo con ReportLab ──
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                        Paragraph, Spacer, Image as RLImage)
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        messagebox.showerror(
            "Librería faltante",
            "Instala reportlab para exportar PDF:\n\npip install reportlab")
        return None

    nombre_sugerido = f"Asistencias_{_safe_filename(nombre_cap)}_{datetime.date.today()}.pdf"
    fname = _pedir_ruta_guardado(nombre_sugerido)
    if fname is None:
        return None

    with sqlite3.connect(DB_PATH) as con:
        rows = con.execute("""
            SELECT p.nombre || ' ' || p.apellido, p.cargo, a.hora_registro
            FROM asistencias a JOIN personas p ON p.id = a.persona_id
            WHERE a.capacitacion_id=? ORDER BY a.hora_registro
        """, (capacitacion_id,)).fetchall()
        cap_info = con.execute(
            "SELECT nombre, descripcion, fecha, firma_responsable, firma_png "
            "FROM capacitaciones WHERE id=?",
            (capacitacion_id,)).fetchone()

    hoy = datetime.date.today().strftime("%d/%m/%Y")
    firma_png_ruta = cap_info[4] if cap_info and len(cap_info) > 4 else None

    AZUL_OSC  = colors.HexColor("#1B2A4A")
    AZUL_MED  = colors.HexColor("#1B4F9C")
    AZUL_CLA  = colors.HexColor("#E8EFF9")
    VERDE_CLA = colors.HexColor("#E8F5EE")
    VERDE     = colors.HexColor("#1A7D4E")
    GRIS_CLA  = colors.HexColor("#F0F2F5")
    TEXTO2    = colors.HexColor("#4A5568")
    BLANCO    = colors.white

    doc = SimpleDocTemplate(fname, pagesize=A4,
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=1.5*cm,  bottomMargin=1.5*cm)
    ancho = A4[0] - 3*cm
    story = []

    st_titulo = ParagraphStyle("titulo", fontName="Helvetica-Bold", fontSize=15,
                                textColor=BLANCO, alignment=TA_CENTER, leading=22)
    st_cap    = ParagraphStyle("cap", fontName="Helvetica-Bold", fontSize=12,
                                textColor=AZUL_MED, alignment=TA_LEFT, leading=18)
    st_meta   = ParagraphStyle("meta", fontName="Helvetica-Oblique", fontSize=8,
                                textColor=TEXTO2, alignment=TA_LEFT, leading=13)
    st_total  = ParagraphStyle("total", fontName="Helvetica-Bold", fontSize=9,
                                textColor=VERDE, alignment=TA_LEFT, leading=14)

    encabezado = [
        [Paragraph("SISTEMA DE ASISTENCIAS A CAPACITACIONES", st_titulo)],
        [Paragraph(f"  {cap_info[0]}", st_cap)],
        [Paragraph(
            f"  Descripción: {cap_info[1] or 'N/A'}   |   "
            f"Fecha: {cap_info[2]}   |   Exportado: {hoy}", st_meta)],
        [Paragraph(f"  Total de asistentes: {len(rows)}", st_total)],
    ]
    enc_t = Table(encabezado, colWidths=[ancho])
    enc_t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (0, 0), AZUL_OSC),
        ("BACKGROUND",    (0, 1), (0, 1), AZUL_CLA),
        ("BACKGROUND",    (0, 2), (0, 2), GRIS_CLA),
        ("BACKGROUND",    (0, 3), (0, 3), VERDE_CLA),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        ("BOX",           (0, 0), (-1, -1), 1, AZUL_MED),
        ("LINEBELOW",     (0, 0), (0, 0),   0.5, AZUL_MED),
        ("LINEBELOW",     (0, 1), (0, 1),   0.5, colors.HexColor("#CBD5E0")),
        ("LINEBELOW",     (0, 2), (0, 2),   0.5, colors.HexColor("#CBD5E0")),
    ]))
    story.append(enc_t)
    story.append(Spacer(1, 0.35 * cm))

    col_ws = [ancho * 0.06, ancho * 0.40, ancho * 0.28, ancho * 0.26]
    datos  = [["N°", "Nombre Completo", "Cargo / Puesto", "Hora de Registro"]]
    for i, (nombre, cargo, hora) in enumerate(rows, 1):
        datos.append([str(i), nombre, cargo or "—", hora])

    tbl = Table(datos, colWidths=col_ws, repeatRows=1)
    ts  = TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  AZUL_MED),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  BLANCO),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0),  9),
        ("ALIGN",         (0, 0), (-1, 0),  "CENTER"),
        ("TOPPADDING",    (0, 0), (-1, 0),  7),
        ("BOTTOMPADDING", (0, 0), (-1, 0),  7),
        ("FONTNAME",      (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE",      (0, 1), (-1, -1), 9),
        ("TOPPADDING",    (0, 1), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 5),
        ("ALIGN",         (0, 1), (0, -1),  "CENTER"),
        ("ALIGN",         (3, 1), (3, -1),  "CENTER"),
        ("ALIGN",         (1, 1), (2, -1),  "LEFT"),
        ("LEFTPADDING",   (1, 1), (2, -1),  6),
        ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E0")),
        ("BOX",           (0, 0), (-1, -1), 1,   AZUL_MED),
    ])
    for i in range(1, len(datos)):
        ts.add("BACKGROUND", (0, i), (-1, i), AZUL_CLA if i % 2 == 0 else BLANCO)
    tbl.setStyle(ts)
    story.append(tbl)
    story.append(Spacer(1, 0.4 * cm))

    # ── Sección de firma ─────────────────────
    firma_val    = cap_info[3] or "—"
    tiene_imagen = firma_png_ruta and os.path.isfile(firma_png_ruta)

    st_firma_lbl = ParagraphStyle("firma_lbl", fontName="Helvetica-Bold", fontSize=9,
                                    textColor=TEXTO2, alignment=TA_LEFT)
    st_firma_nom = ParagraphStyle("firma_nom", fontName="Helvetica-Bold", fontSize=9,
                                    textColor=AZUL_MED, alignment=TA_CENTER)

    if tiene_imagen:
        img_firma = RLImage(firma_png_ruta, width=3.5 * cm, height=1.5 * cm)
        firma_data     = [[Paragraph("Responsable de la Capacitación:", st_firma_lbl),
                            Paragraph(firma_val, st_firma_nom),
                            img_firma]]
        col_ws_firma   = [ancho * 0.35, ancho * 0.35, ancho * 0.30]
        estilo_extra   = [("BACKGROUND", (2, 0), (2, 0), BLANCO),
                            ("ALIGN",      (2, 0), (2, 0), "CENTER"),
                            ("VALIGN",     (2, 0), (2, 0), "MIDDLE")]
    else:
        firma_data     = [[Paragraph("Responsable de la Capacitación:", st_firma_lbl),
                            Paragraph(firma_val, st_firma_nom)]]
        col_ws_firma   = [ancho * 0.45, ancho * 0.55]
        estilo_extra   = []

    firma_t = Table(firma_data, colWidths=col_ws_firma)
    firma_style = TableStyle([
        ("BACKGROUND",    (0, 0), (0, 0), GRIS_CLA),
        ("BACKGROUND",    (1, 0), (1, 0), AZUL_CLA),
        ("FONTSIZE",      (0, 0), (-1, -1), 9),
        ("ALIGN",         (0, 0), (0, 0),  "LEFT"),
        ("ALIGN",         (1, 0), (1, 0),  "CENTER"),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING",   (0, 0), (0, 0),  10),
        ("BOX",           (0, 0), (-1, -1), 0.5, AZUL_MED),
        ("INNERGRID",     (0, 0), (-1, -1), 0.5, AZUL_MED),
    ])
    for cmd in estilo_extra:
        firma_style.add(*cmd)
    firma_t.setStyle(firma_style)
    story.append(firma_t)
    story.append(Spacer(1, 0.3 * cm))

    st_pie = ParagraphStyle("pie", fontName="Helvetica-Oblique", fontSize=7,
                            textColor=colors.HexColor("#A0AEC0"), alignment=TA_CENTER)
    pie_t = Table([[Paragraph("— HOPD —", st_pie)]], colWidths=[ancho])
    pie_t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), GRIS_CLA),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(pie_t)

    doc.build(story)
    return fname